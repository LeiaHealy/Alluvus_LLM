import sys
import os
from docx import Document
from docx.shared import Pt
from docx.shared import Inches, Cm
from docx.enum.style import WD_STYLE_TYPE as ST

doc = Document()

styles = doc.styles

heading_style = styles.add_style("heading", ST.CHARACTER)
heading_style.font.name = 'Arial'
heading_style.font.size = Pt(15)

heading_txt = 'This is the main heading'
heading = doc.add_paragraph()
heading_runner = heading.add_run(heading_txt, heading_style)
heading_runner.bold = True
heading.alignment = 1


subheading_txt = 'This is the sub-heading'
subheading = doc.add_paragraph()
subheading_runner = subheading.add_run(subheading_txt)
subheading_runner.italic = True
subheading.alignment = 1

paragraph = doc.add_paragraph('testing paragraph goes here. Ya Know I have always kinda wanted one of these. Just a place to randomly babble nonstop into the void. It could be nice but idk. Ya know?')
paragraph = doc.add_paragraph('This is the second paragraph. This is the second paragraph. This is the second paragraph. This is the second paragraph. This is the second paragraph. This is the second paragraph')
   
# body_style = styles.add_style("body", ST.CHARACTER)
# body_style.font.name = 'Arial'
# body_style.font.size = Pt(12)
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
    

filename = 'test_style.docx'
doc.save('/Users/erika/Documents/Alluvus/GPT_FineTuning/tests/smpl_out/'+filename)