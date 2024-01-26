import os
import sys
import json
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE as ST

from openai import OpenAI

client = OpenAI()


def convert_to_docx(filename, headline_txt, subheadline_txt, data_str, pr_type):
    # Get the directory of the executable
    executable_dir = os.path.dirname(sys.executable)

    # Path for the output directory
    output_dir = os.path.join(executable_dir, 'output')

    # Check and create the output directory
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    print(f"Output directory exists: {os.path.exists(output_dir)} at {output_dir}")

    # print(f'data_str\nType: {type(data_str)}\n{data_str}\n================')
    
    # Load the JSON string into a Python dictionary
    data = json.loads(data_str)

    # print(f'json_data\nType: {type(data)}\n{data}\n================')

    # TODO: There's a good chance this WILL be an issue in the future. Find more permanent fix for single line thing
    # It is indeed an issue
    if pr_type == 0:
        # Extract the 'press_release' content and replace escape sequences
        press_release = data["press_release"].replace("\\n\\n", "\n")
    elif pr_type == 1:
        press_release = data["consolidated_doc"].replace("\\n\\n", "\n")
        
    # Create a new Word document
    doc = Document()
    
    styles = doc.styles

    headline_style = styles.add_style("heading", ST.CHARACTER)
    headline_style.font.name = 'Arial'
    headline_style.font.size = Pt(15)
    headline = doc.add_paragraph()

    headline_runner = headline.add_run(headline_txt, headline_style)
    headline_runner.bold = True
    headline.alignment = 1

    subheadline = doc.add_paragraph()
    subheadline_runner = subheadline.add_run(subheadline_txt)
    subheadline_runner.italic = True
    subheadline.alignment = 1


    # Define style for content
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    split_pr = press_release.split('\n')

    # print(f'press release.split(newline)\n================\n{split_pr}')
    
    # Add the content to the document. Split by '\n' and add each line as a new paragraph
    for line in press_release.split('\n'):
        # print(f"{line} : {len(line)}")
        if len(line) > 0:
            doc.add_paragraph(line)
    

    # Save the document   
    doc_path = output_dir + '/' + filename + '.docx'
    doc.save(doc_path)

def gpt_press_release(filename, headline, subheadline, announcement, quotes, info_text, boilerplate):
    std_press_release_func = {
            "name": "gen_press_release",
            "description": "Generate a standard press release.",
            "parameters": {
                "type": "object",
                "properties": {
                    "press_release": {
                        "type": "string",
                        "description": """The press release generated from the info given. Have FOR INTERNAL USE ONLY be at the top.
                                        Be professional. Be succinct, but make it cohesive and add filler to extend certain aspects.
                                        Only add ONE new line between paragraphs. If a word is surrounded by '_', it should be italicized
                                        e.g. _italicize_ . If a word is surrounded by '*' it should be bolded. e.g. *bold*""",
                    },
                },
                "required": ["press_release"],
            }
        }
    
    # info = ""
    # for ln in info_array:
    #     info += '\n' + str(ln)

    if len(quotes) > 1:
        content = f"""Generate a press release regarding {announcement}. 
        Additional info to include: {info_text}.
        Generate quotes surrounded by "" using information from here: {quotes}.
        Add boilerplate: {boilerplate}"""
    else:
        content = f"""Generate a press release regarding {announcement}. 
        Additional info to include: {info_text}.
        Add boilerplate: {boilerplate}"""

    # print(f'Announcement\n############################{announcement}\nType: {type(announcement)}\n') 
    # print(f'info\n############################{info_text}\nType: {type(info_text)}\n')
    # print(f'boilerplate\n############################{boilerplate}\nType: {type(boilerplate)}\n')

    r = client.chat.completions.create(model="gpt-4",
    temperature=0.0,
    messages=[{"role": "user", "content": content}],
    functions=[std_press_release_func])

    output_str = r.choices[0].message.function_call.arguments

    convert_to_docx(filename, headline, subheadline, output_str, 0)
    print(f'Finished generating {filename}')

def gpt_media_advisory(filename, headline, subheadline, announcement, what, where, when, event_overview, boilerplate):
    print(
    f'INITIAL INPUT\n'
    f'Filename: {filename}\n'
    f'announcement: {announcement}\n'
    f'what: {what}\n'
    f'where: {where}\n'
    f'when: {when}\n'
    f'Event Overview: {event_overview}\n'
    f'Boilerplate: {boilerplate}'
    )


    media_advisory_func = {
        "name": "gen_media_advisory",
        "description": "Generate a document consolidating the given media advisory information",
        "parameters": {
            "type": "object",
            "properties": {
                "consolidated_doc": {
                    "type": "string",
                    "description": """You are generating a media advisory. It is similar to a press release,
                    but in a relatively different format. Media Advisory is bold and the title. Underneath, also bolded, is the announcement to be made. 
                    There will then be a "What", "Where", and "When" section, where each will have their own section. 
                    e.g. What:/t [what is happening]. "When" can have multiple sections. e.g. one event that is 
                    from 8:00am to 6:00pm, or it can have multiple, overlapping events. Underneath all of this, 
                    there is the event overview section. Finally, tack on the boilerplate at the end.""",
                },
            },
            "required": ["consolidated_doc"],
        }
    }

    # for ln in when:
    #     when_info += '\n' + str(ln)
    content = f"""Generate a media advisory announcing {announcement}.
    what:{what}, where:{where}, when:{when}.
    Event overview: {event_overview}
    Boilerplate: {boilerplate}"""

    print(f"================\nCOMBINED CONTENT\n{content}\n================")

    r = client.chat.completions.create(model="gpt-4",
    temperature=0.0,
    messages=[{"role": "user", "content": content}],
    functions=[media_advisory_func])
    output_str = r.choices[0].message.function_call.arguments
    print(f"AI OUTPUT\n#################################\n {r}\n\n#################################\nOUTPUT STRING: \n#################################\n{output_str}\n######################")
    convert_to_docx(filename, headline, subheadline, output_str, 1)
    print(f'Finished generating {filename}')